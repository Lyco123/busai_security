# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("agent/docs/BUS-Agent-ops-manual.docx")


def set_font(run, name="宋体", size=10.5, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    for line in text.strip().split("\n"):
        run = p.add_run(line + "\n")
        set_font(run, "Consolas", 9)
    return p


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, True)
        shade(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    for _ in range(6):
        doc.add_paragraph()
    for text in ["BUS Agent 后端服务", "运行维护手册"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=24, bold=True)
    for _ in range(8):
        doc.add_paragraph()
    for text in ["文档版本：V1.0", "适用范围：agent-deployment 本地化部署版本", "编制日期：2026年7月9日"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, size=12)
    doc.add_page_break()

    add_heading(doc, "目 录", 1)
    for item in [
        "第一章 总体情况",
        "1.1 编写目的",
        "1.2 建设背景",
        "1.3 功能概述",
        "1.4 系统部署架构",
        "第二章 系统运行环境",
        "2.1 服务器信息",
        "2.2 软件环境",
        "2.3 应用部署目录",
        "2.4 网络与端口",
        "2.5 外部依赖",
        "第三章 系统部署说明",
        "3.1 首次部署准备",
        "3.2 获取代码",
        "3.3 配置环境变量",
        "3.4 构建并启动服务",
        "3.5 健康检查",
        "3.6 停止、重启与更新",
        "第四章 系统维护说明",
        "4.1 维护时间",
        "4.2 维护准备工作",
        "4.3 日常巡检内容",
        "4.4 Docker 容器维护",
        "4.5 SQLite 数据库维护",
        "4.6 日志维护",
        "4.7 数据备份与恢复",
        "4.8 版本更新与回滚",
        "4.9 安全维护",
        "第五章 常见故障处理",
        "5.1 服务无法启动",
        "5.2 Docker 构建失败",
        "5.3 端口无法访问",
        "5.4 健康检查失败",
        "5.5 SQLite 数据库异常",
        "5.6 模型接口异常",
        "5.7 MCP 接口异常",
        "5.8 磁盘空间不足",
        "5.9 VPN/堡垒机连接异常",
        "第六章 附录",
        "6.1 常用命令",
        "6.2 环境变量清单",
        "6.3 目录结构说明",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_font(run)
    doc.add_page_break()

    add_heading(doc, "第一章 总体情况", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_paragraph(
        doc,
        "本手册用于规范 BUS Agent 后端服务本地化部署后的运行维护工作，指导运维人员完成系统巡检、服务启停、日志查看、数据备份、故障排查、版本更新和回滚等操作，保障系统稳定、可靠、安全运行。",
    )
    add_heading(doc, "1.2 建设背景", 2)
    add_paragraph(
        doc,
        "BUS Agent 后端服务按本地化部署场景设计，采用 Node.js 服务、Docker Compose 编排和 SQLite 本地文件数据库的运行形态。系统通过 agent-deployment 目录统一管理启动入口、运行配置、数据库初始化、容器构建和运维脚本，满足内网服务器部署、VPN/堡垒机访问、数据本地留存和日常运维管理要求。",
    )
    add_heading(doc, "1.3 功能概述", 2)
    add_paragraph(
        doc,
        "BUS Agent 后端服务主要提供 AI 助手会话、结构化报告生成、规则配置、研究评估、会话管理、实体别名管理、MCP 工具调用、知识库检索代理等能力。系统对外提供 /api/agent/* 接口，由前端或其他业务系统调用。",
    )
    add_heading(doc, "1.4 系统部署架构", 2)
    add_paragraph(
        doc,
        "系统采用单机 Docker Compose 部署。agent-deployment 作为部署壳，复用 agent 目录中的业务代码；Node.js 容器运行后端服务；SQLite 数据文件通过 Docker volume 持久化保存；如需公网或统一入口访问，可在服务器前置 Nginx 反向代理。",
    )
    add_table(
        doc,
        ["组件", "说明"],
        [
            ["agent", "业务代码目录，包含接口路由、领域服务、工具调用逻辑和数据库 migrations。"],
            ["agent-deployment", "本地化部署目录，包含 Node 启动入口、SQLite 数据访问适配器、Dockerfile、docker-compose.yml 和 .env.example。"],
            ["bus-agent 容器", "Node.js 后端服务容器，默认监听 3000 端口。"],
            ["SQLite", "本地文件数据库，容器内默认路径为 /app/data/bus-agent.sqlite。"],
            ["Docker volume", "持久化保存 SQLite 数据文件，默认 volume 名称为 bus_agent_data。"],
        ],
    )

    add_heading(doc, "第二章 系统运行环境", 1)
    add_heading(doc, "2.1 服务器信息", 2)
    add_table(
        doc,
        ["项目", "当前值", "备注"],
        [
            ["服务器 IP", "待补充", "按实际部署服务器填写。"],
            ["主机名", "dx92108", "根据 WebCLI 当前提示记录。"],
            ["操作系统", "openEuler", "服务器已确认为 openEuler 系统。"],
            ["登录方式", "VPN + 堡垒机 WebCLI", "通过堡垒机系统进入服务器命令行。"],
            ["操作系统用户", "root", "当前部署操作使用 root 用户。"],
        ],
    )
    add_heading(doc, "2.2 软件环境", 2)
    add_table(
        doc,
        ["软件组件", "版本/命令", "说明"],
        [
            ["Git", "git", "用于拉取 GitHub 仓库代码。"],
            ["Docker", "18.09.0 EulerVersion 18.09.0.353", "已在 openEuler 上安装并启动。"],
            ["docker-compose", "1.22.0", "使用老版 docker-compose 命令，不使用 docker compose 插件命令。"],
            ["Node.js", "node:22-bookworm-slim 镜像", "由 Dockerfile 在容器内提供。"],
            ["SQLite", "better-sqlite3 内嵌访问", "不单独部署数据库服务。"],
        ],
    )
    add_heading(doc, "2.3 应用部署目录", 2)
    add_table(
        doc,
        ["目录/文件", "说明"],
        [
            ["/opt/BUS", "代码仓库根目录。"],
            ["/opt/BUS/agent", "业务代码目录。"],
            ["/opt/BUS/agent-deployment", "部署目录，执行 docker-compose 命令的位置。"],
            ["/opt/BUS/agent-deployment/.env", "生产环境配置文件，不提交 Git。"],
            ["/app/data/bus-agent.sqlite", "容器内 SQLite 数据库文件路径。"],
            ["bus_agent_data", "Docker volume 名称，用于持久化 SQLite 数据。"],
        ],
    )
    add_heading(doc, "2.4 网络与端口", 2)
    add_table(
        doc,
        ["访问项", "地址/端口", "说明"],
        [
            ["后端服务", "127.0.0.1:3000", "容器映射到宿主机 3000 端口。"],
            ["健康检查", "http://127.0.0.1:3000/api/agent/health", "用于判断后端服务是否正常。"],
            ["业务接口", "/api/agent/*", "前端或上游系统调用路径。"],
            ["Nginx 反向代理", "待补充", "如部署 Nginx，填写域名和监听端口。"],
        ],
    )
    add_heading(doc, "2.5 外部依赖", 2)
    add_table(
        doc,
        ["依赖项", "环境变量", "说明"],
        [
            ["大模型服务", "OPENAI_BASE_URL、OPENAI_API_KEY、OPENAI_MODEL", "用于普通对话、路由、规则配置等模型调用。"],
            ["Embedding 服务", "OPENAI_EMBEDDING_MODEL", "用于规则/场景向量生成。"],
            ["MCP 服务", "MCP_SERVER_URL、MCP_ACCESS_TOKEN", "用于画像数据查询。"],
            ["知识库服务", "KB_API_BASE_URL、KB_DEFAULT_ID、KB_TOOL_ENABLED", "如启用知识库检索工具，需要配置。"],
        ],
    )

    add_heading(doc, "第三章 系统部署说明", 1)
    add_heading(doc, "3.1 首次部署准备", 2)
    add_paragraph(doc, "首次部署前需要确认服务器可通过 VPN/堡垒机登录，Docker 服务已启动，docker-compose 命令可用，服务器可访问 GitHub 或可通过堡垒机上传代码包。")
    add_code(doc, "docker version\ndocker-compose version\nsystemctl status docker")
    add_heading(doc, "3.2 获取代码", 2)
    add_paragraph(doc, "推荐直接在服务器上通过 Git 拉取完整仓库。agent-deployment 依赖 agent 目录中的业务代码，因此不能只上传 agent-deployment 目录。")
    add_code(doc, "cd /opt\ngit clone https://github.com/forceve/BUS.git\ncd /opt/BUS/agent-deployment")
    add_heading(doc, "3.3 配置环境变量", 2)
    add_paragraph(doc, "复制环境变量模板并按实际环境修改。生产环境 .env 文件包含密钥，不应提交到 Git。")
    add_code(doc, "cd /opt/BUS/agent-deployment\ncp .env.example .env\nvi .env")
    add_paragraph(doc, "至少需要配置 OPENAI_API_KEY、OPENAI_BASE_URL、OPENAI_MODEL、OPENAI_ROUTER_MODEL、OPENAI_WORKER_MODEL、OPENAI_EMBEDDING_MODEL、MCP_SERVER_URL、MCP_ACCESS_TOKEN。")
    add_heading(doc, "3.4 构建并启动服务", 2)
    add_code(doc, "cd /opt/BUS/agent-deployment\ndocker-compose up -d --build")
    add_paragraph(doc, "容器启动时会自动执行 SQLite migrations，然后启动 Node.js 后端服务。")
    add_heading(doc, "3.5 健康检查", 2)
    add_code(doc, "curl http://127.0.0.1:3000/api/agent/health")
    add_paragraph(doc, "正常情况下返回 JSON，其中 status 字段为 ok。")
    add_code(doc, '{"status":"ok","timestamp":"2026-07-09T00:00:00.000Z"}')
    add_heading(doc, "3.6 停止、重启与更新", 2)
    add_code(doc, "cd /opt/BUS/agent-deployment\ndocker-compose ps\ndocker-compose logs -f bus-agent\ndocker-compose restart bus-agent\ndocker-compose down")

    add_heading(doc, "第四章 系统维护说明", 1)
    maintenance = [
        ("4.1 维护时间", "建议每周进行一次例行巡检；版本更新、数据恢复、系统重启等操作应安排在业务低峰期进行。紧急故障处理不受维护窗口限制，但应保留操作记录。"),
        ("4.2 维护准备工作", "维护人员需要具备 VPN/堡垒机登录权限、服务器 root 或等效操作权限、GitHub 仓库访问权限、生产 .env 配置查看或更新权限，以及必要的模型服务和 MCP 服务访问确认能力。"),
    ]
    for title, text in maintenance:
        add_heading(doc, title, 2)
        add_paragraph(doc, text)
    add_heading(doc, "4.3 日常巡检内容", 2)
    add_table(
        doc,
        ["巡检项", "命令/方式", "正常标准"],
        [
            ["系统负载", "top 或 uptime", "CPU、Load 无长期异常升高。"],
            ["内存使用", "free -h", "可用内存充足，无持续 OOM。"],
            ["磁盘空间", "df -h", "数据盘和 Docker 目录使用率建议低于 80%。"],
            ["Docker 服务", "systemctl status docker", "服务状态为 active/running。"],
            ["容器状态", "docker-compose ps", "bus-agent 状态为 Up。"],
            ["应用健康", "curl /api/agent/health", "返回 status=ok。"],
            ["应用日志", "docker-compose logs --tail=200 bus-agent", "无持续报错。"],
        ],
    )
    add_heading(doc, "4.4 Docker 容器维护", 2)
    add_code(doc, "cd /opt/BUS/agent-deployment\ndocker-compose ps\ndocker-compose logs --tail=200 bus-agent\ndocker-compose logs -f bus-agent\ndocker-compose restart bus-agent")
    add_paragraph(doc, "如需清理无用镜像，应先确认当前容器运行正常，不得删除 bus_agent_data volume。")
    add_heading(doc, "4.5 SQLite 数据库维护", 2)
    add_paragraph(doc, "SQLite 数据库不作为独立服务运行，其数据文件由 bus-agent 容器访问，并通过 Docker volume 持久化保存。维护时重点关注数据文件存在性、文件大小、磁盘空间、备份文件完整性。")
    add_code(doc, "cd /opt/BUS/agent-deployment\ndocker-compose exec bus-agent ls -lh /app/data\ndocker volume ls")
    add_heading(doc, "4.6 日志维护", 2)
    add_paragraph(doc, "应用日志通过 Docker 标准输出查看。日常巡检保留关键错误信息；发生故障时应记录故障时间、请求路径、错误堆栈、环境变量变更和发布版本。")
    add_code(doc, "docker-compose logs --tail=500 bus-agent\ndocker-compose logs -f bus-agent")
    add_heading(doc, "4.7 数据备份与恢复", 2)
    add_paragraph(doc, "建议在每次版本更新前和每日低峰期备份 SQLite 数据文件。备份文件可保存在 Docker volume 内，也可复制到宿主机独立备份目录。")
    add_code(doc, 'cd /opt/BUS/agent-deployment\ndocker-compose exec bus-agent sh -c "cp /app/data/bus-agent.sqlite /app/data/bus-agent.$(date +%F-%H%M%S).bak"\ndocker-compose exec bus-agent ls -lh /app/data')
    add_heading(doc, "4.8 版本更新与回滚", 2)
    add_code(doc, "cd /opt/BUS\ngit pull\ncd /opt/BUS/agent-deployment\ndocker-compose up -d --build\ndocker-compose logs -f bus-agent")
    add_paragraph(doc, "如新版本异常，可通过 git checkout 回退到上一稳定提交并重新构建。")
    add_heading(doc, "4.9 安全维护", 2)
    add_paragraph(doc, "生产 .env 文件应限制访问权限，不得上传到代码仓库。MCP_ACCESS_TOKEN、OPENAI_API_KEY 等密钥应定期轮换。服务器仅开放必要端口，后台接口如需公网访问，应通过 Nginx、VPN 或统一鉴权入口进行访问控制。")

    add_heading(doc, "第五章 常见故障处理", 1)
    for title, text in [
        ("5.1 服务无法启动", "查看 docker-compose logs -f bus-agent，重点检查 .env 是否缺少模型或 MCP 配置，SQLite 数据目录是否可写，容器是否构建成功。"),
        ("5.2 Docker 构建失败", "确认 Docker 服务运行，网络可拉取 node:22-bookworm-slim 镜像；如 GitHub 或 npm 访问受限，需要配置代理或离线镜像。"),
        ("5.3 端口无法访问", "执行 docker-compose ps、ss -lntp | grep 3000，确认容器状态和端口映射；如有防火墙或安全组，需要放通访问端口。"),
        ("5.4 健康检查失败", "执行 curl http://127.0.0.1:3000/api/agent/health，并查看应用日志。若容器不断重启，优先检查 migrations 和 .env。"),
        ("5.5 SQLite 数据库异常", "检查 /app/data/bus-agent.sqlite 是否存在，磁盘是否写满，Docker volume 是否误删，必要时从最近备份恢复。"),
        ("5.6 模型接口异常", "检查 OPENAI_API_KEY、OPENAI_BASE_URL、OPENAI_MODEL 等配置，确认服务器可访问模型服务地址。"),
        ("5.7 MCP 接口异常", "检查 MCP_SERVER_URL、MCP_ACCESS_TOKEN、CF_ACCESS_CLIENT_ID、CF_ACCESS_CLIENT_SECRET、MCP_REQUEST_TIMEOUT_MS，并确认服务器网络可达 MCP 服务。"),
        ("5.8 磁盘空间不足", "执行 df -h、du -h -x --max-depth=1 /var/lib/docker，清理无用镜像和旧备份，但不得删除 bus_agent_data volume。"),
        ("5.9 VPN/堡垒机连接异常", "确认 VPN 账号、堡垒机权限、服务器授权列表和 WebCLI 会话是否正常；必要时联系平台管理员开通或恢复权限。"),
    ]:
        add_heading(doc, title, 2)
        add_paragraph(doc, text)

    add_heading(doc, "第六章 附录", 1)
    add_heading(doc, "6.1 常用命令", 2)
    add_code(doc, "cd /opt/BUS/agent-deployment\ndocker-compose up -d --build\ndocker-compose ps\ndocker-compose logs -f bus-agent\ndocker-compose restart bus-agent\ndocker-compose down\ncurl http://127.0.0.1:3000/api/agent/health\ndf -h\nfree -h\ntop")
    add_heading(doc, "6.2 环境变量清单", 2)
    add_table(
        doc,
        ["环境变量", "是否必填", "说明"],
        [
            ["PORT", "否", "服务端口，默认 3000。"],
            ["HOST", "否", "监听地址，默认 0.0.0.0。"],
            ["SQLITE_DB_PATH", "是", "SQLite 数据文件路径，容器内默认 /app/data/bus-agent.sqlite。"],
            ["OPENAI_API_KEY", "是", "模型服务 API Key。"],
            ["OPENAI_BASE_URL", "是", "模型服务基础地址。"],
            ["OPENAI_MODEL", "是", "默认模型。"],
            ["OPENAI_ROUTER_MODEL", "建议", "路由模型。"],
            ["OPENAI_WORKER_MODEL", "建议", "专家/报告执行模型。"],
            ["OPENAI_EMBEDDING_MODEL", "建议", "Embedding 模型。"],
            ["OPENAI_LOCAL_BASE_URL", "否", "本地模型兜底地址。"],
            ["OPENAI_LOCAL_MODEL", "否", "本地模型名称。"],
            ["OPENAI_REPORT_BASE_URL / OPENAI_REPORT_URL", "否", "报告生成专用模型地址。"],
            ["OPENAI_REPORT_API_KEY", "否", "报告生成专用模型 Key。"],
            ["MCP_SERVER_URL", "是", "MCP 服务地址。"],
            ["MCP_ACCESS_TOKEN", "是", "MCP 访问令牌。"],
            ["MCP_REQUEST_TIMEOUT_MS", "否", "MCP 请求超时时间。"],
            ["KB_API_BASE_URL", "否", "知识库服务地址。"],
            ["KB_TOOL_ENABLED", "否", "是否启用知识库工具。"],
            ["KB_DEFAULT_ID", "否", "默认知识库 ID。"],
            ["CORS_ALLOWED_ORIGINS", "否", "允许跨域来源列表。"],
            ["OUTPUT_FORMAT", "否", "输出格式，默认 markdown。"],
        ],
    )
    add_heading(doc, "6.3 目录结构说明", 2)
    add_code(doc, "/opt/BUS\n  agent/\n    src/\n    migrations/\n  agent-deployment/\n    Dockerfile\n    docker-compose.yml\n    .env\n    src/\n      node-server.ts\n      sqlite-adapter.ts\n      migrate.ts")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT.resolve())
